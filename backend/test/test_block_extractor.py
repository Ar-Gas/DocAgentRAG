import os
import subprocess
import sys
from pathlib import Path

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

from utils.block_extractor import (
    BLOCK_INDEX_VERSION,
    TABLE_SPLIT_ROW_THRESHOLD,
    _extract_heading_level,
    _split_table_block,
    assign_block_ids,
    compute_indexed_content_hash,
    extract_structured_blocks,
)


def test_indexed_content_hash_stable_for_semantically_identical_blocks():
    blocks_a = [
        {
            "block_type": "paragraph",
            "heading_path": ["  第一章\n总则  ", "  适用\t 范围 \n说明"],
            "page_number": 1,
            "text": "第一条  目的\r\n\r\n\r\n适用于全体员工  \r\n",
        }
    ]
    blocks_b = [
        {
            "block_type": "paragraph",
            "heading_path": ["第一章 总则", "适用 范围 说明"],
            "page_number": 1,
            "text": "第一条  目的\n\n适用于全体员工\n",
        }
    ]

    assert compute_indexed_content_hash(blocks_a) == compute_indexed_content_hash(blocks_b)


def test_indexed_content_hash_changes_when_internal_text_spacing_changes():
    blocks_a = [
        {
            "block_type": "paragraph",
            "heading_path": ["第一章 总则"],
            "page_number": 1,
            "text": "第一条  目的",
        }
    ]
    blocks_b = [
        {
            "block_type": "paragraph",
            "heading_path": ["第一章 总则"],
            "page_number": 1,
            "text": "第一条 目的",
        }
    ]

    assert compute_indexed_content_hash(blocks_a) != compute_indexed_content_hash(blocks_b)


def test_indexed_content_hash_changes_when_page_number_changes():
    blocks_a = [
        {
            "block_type": "paragraph",
            "heading_path": ["第一章 总则"],
            "page_number": 1,
            "text": "相同内容",
        }
    ]
    blocks_b = [
        {
            "block_type": "paragraph",
            "heading_path": ["第一章 总则"],
            "page_number": 2,
            "text": "相同内容",
        }
    ]

    assert compute_indexed_content_hash(blocks_a) != compute_indexed_content_hash(blocks_b)


def test_assign_block_ids_deterministic_with_document_id_version_and_index():
    blocks = [
        {"block_type": "paragraph", "heading_path": [], "text": "A"},
        {"block_type": "paragraph", "heading_path": ["H1"], "text": "B"},
    ]

    assign_block_ids(blocks, document_id="doc-123", index_version=BLOCK_INDEX_VERSION)

    assert blocks[0]["block_index"] == 0
    assert blocks[1]["block_index"] == 1
    assert blocks[0]["block_id"] == "doc-123:block-v1:0"
    assert blocks[1]["block_id"] == "doc-123:block-v1:1"


def test_extract_heading_level_supports_english_and_chinese_heading_styles():
    assert _extract_heading_level("Heading 2") == 2
    assert _extract_heading_level("标题 3") == 3
    assert _extract_heading_level("Normal") is None


def test_split_table_block_uses_row_threshold_and_repeats_header():
    rows = ["列1 | 列2"] + [f"行{i} | 值{i}" for i in range(1, TABLE_SPLIT_ROW_THRESHOLD + 2)]

    blocks = _split_table_block(rows, heading_path=["第一章 总则"], page_number=3)

    assert len(blocks) == 2
    assert all(block["block_type"] == "table" for block in blocks)
    assert all(block["heading_path"] == ["第一章 总则"] for block in blocks)
    assert all(block["page_number"] == 3 for block in blocks)
    assert all(block["text"].split("\n")[0] == "列1 | 列2" for block in blocks)
    assert len(blocks[0]["text"].split("\n")) == TABLE_SPLIT_ROW_THRESHOLD


def test_extract_structured_blocks_smoke_docx_has_required_fields():
    sample_docx = Path(__file__).parent / "test_date" / "sample.docx"

    payload = extract_structured_blocks(str(sample_docx), document_id="sample-docx")

    assert payload["index_version"] == BLOCK_INDEX_VERSION
    assert payload["indexed_content_hash"]
    assert payload["blocks"]
    for index, block in enumerate(payload["blocks"]):
        assert block["block_id"] == f"sample-docx:{BLOCK_INDEX_VERSION}:{index}"
        assert block["block_type"]
        assert "heading_path" in block
        assert "page_number" in block


def test_extract_structured_blocks_docx_with_table_emits_table_block(tmp_path: Path):
    sample_docx = tmp_path / "table.docx"
    document = Document()
    document.add_heading("第一章 总则", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "方向"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "人工智能"
    document.save(sample_docx)

    runner = subprocess.run(
        [
            sys.executable,
            "-c",
            (
                "import json, sys; "
                "sys.path.insert(0, %r); "
                "from utils.block_extractor import extract_structured_blocks; "
                "payload = extract_structured_blocks(%r, document_id='table-docx'); "
                "assert any(block['block_type'] == 'table' for block in payload['blocks']); "
                "assert any(block['heading_path'] == ['第一章 总则'] for block in payload['blocks']); "
                "print(json.dumps({'ok': True}, ensure_ascii=False))"
            )
            % (os.path.dirname(os.path.dirname(os.path.abspath(__file__))), str(sample_docx)),
        ],
        capture_output=True,
        text=True,
    )

    assert runner.returncode == 0, runner.stderr or runner.stdout


def test_extract_structured_blocks_smoke_pdf_has_page_numbers():
    sample_pdf = Path(__file__).parent / "test_date" / "test_pdf_1.pdf"

    payload = extract_structured_blocks(str(sample_pdf), document_id="sample-pdf")

    assert payload["index_version"] == BLOCK_INDEX_VERSION
    assert payload["blocks"]
    assert all(block["page_number"] is not None for block in payload["blocks"])
